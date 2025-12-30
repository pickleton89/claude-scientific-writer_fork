"""
Oligon Brand Adapter for python-docx (Word Documents)

This module provides functions and styles for creating Word documents
with Oligon scientific brand identity using python-docx.

Usage:
    from oligon_brand.adapters.docx_adapter import (
        create_brand_document,
        apply_brand_styles,
        BRAND_COLORS,
    )

    doc = create_brand_document()
    doc.add_heading("Report Title", level=1)
    doc.add_paragraph("Body text here...")
    doc.save("output.docx")
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional, Tuple

# =============================================================================
# BRAND COLORS (RGB tuples)
# =============================================================================

BRAND_COLORS = {
    # Primary highlight
    'brand_blue': RGBColor(0x2D, 0xB2, 0xE8),

    # Neutrals
    'dark_gray': RGBColor(0x22, 0x22, 0x22),
    'medium_gray': RGBColor(0x66, 0x66, 0x66),
    'muted_gray': RGBColor(0x99, 0x99, 0x99),
    'light_gray': RGBColor(0xBD, 0xBD, 0xBD),

    # Contrast
    'contrast_orange': RGBColor(0xE8, 0x62, 0x2D),
    'dark_warm': RGBColor(0x7D, 0x25, 0x0F),

    # Supporting blues
    'medium_blue': RGBColor(0x15, 0x8B, 0xBB),
    'dark_teal': RGBColor(0x0F, 0x5D, 0x7D),

    # Structure
    'black': RGBColor(0x00, 0x00, 0x00),
    'white': RGBColor(0xFF, 0xFF, 0xFF),
    'gridline': RGBColor(0xE5, 0xE5, 0xE5),
}

# Hex strings for reference
BRAND_HEX = {
    'brand_blue': '#2DB2E8',
    'dark_gray': '#222222',
    'medium_gray': '#666666',
    'muted_gray': '#999999',
    'light_gray': '#BDBDBD',
    'contrast_orange': '#E8622D',
    'dark_warm': '#7D250F',
    'medium_blue': '#158BBB',
    'dark_teal': '#0F5D7D',
    'black': '#000000',
    'white': '#FFFFFF',
    'gridline': '#E5E5E5',
}

# =============================================================================
# TYPOGRAPHY SETTINGS
# =============================================================================

FONT_NAME = 'Arial'

FONT_SIZES = {
    'body': Pt(11),
    'heading1': Pt(14),
    'heading2': Pt(12),
    'heading3': Pt(11),
    'caption': Pt(10),
    'footnote': Pt(9),
    'small': Pt(8),
}

LINE_SPACING = {
    'body': 1.15,
    'heading': 1.2,
}

# =============================================================================
# DOCUMENT CREATION
# =============================================================================

def create_brand_document() -> Document:
    """
    Create a new Word document with brand-compliant styles applied.

    Returns
    -------
    Document
        A python-docx Document with brand styles defined

    Examples
    --------
    >>> doc = create_brand_document()
    >>> doc.add_heading("Introduction", level=1)
    >>> doc.add_paragraph("This is body text.")
    >>> doc.save("report.docx")
    """
    doc = Document()
    apply_brand_styles(doc)
    return doc


def apply_brand_styles(doc: Document) -> None:
    """
    Apply Oligon brand styles to an existing document.

    Parameters
    ----------
    doc : Document
        The python-docx Document to style
    """
    styles = doc.styles

    # --- Normal (Body) Style ---
    style_normal = styles['Normal']
    font = style_normal.font
    font.name = FONT_NAME
    font.size = FONT_SIZES['body']
    font.color.rgb = BRAND_COLORS['black']

    paragraph_format = style_normal.paragraph_format
    paragraph_format.line_spacing = LINE_SPACING['body']
    paragraph_format.space_after = Pt(6)

    # Set East Asian font (for proper Arial rendering)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # --- Heading 1 ---
    style_h1 = styles['Heading 1']
    font = style_h1.font
    font.name = FONT_NAME
    font.size = FONT_SIZES['heading1']
    font.bold = True
    font.color.rgb = BRAND_COLORS['black']

    paragraph_format = style_h1.paragraph_format
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_after = Pt(6)

    # --- Heading 2 ---
    style_h2 = styles['Heading 2']
    font = style_h2.font
    font.name = FONT_NAME
    font.size = FONT_SIZES['heading2']
    font.bold = True
    font.color.rgb = BRAND_COLORS['black']

    paragraph_format = style_h2.paragraph_format
    paragraph_format.space_before = Pt(10)
    paragraph_format.space_after = Pt(4)

    # --- Heading 3 ---
    style_h3 = styles['Heading 3']
    font = style_h3.font
    font.name = FONT_NAME
    font.size = FONT_SIZES['heading3']
    font.bold = True
    font.color.rgb = BRAND_COLORS['black']

    paragraph_format = style_h3.paragraph_format
    paragraph_format.space_before = Pt(8)
    paragraph_format.space_after = Pt(4)

    # --- Caption Style ---
    try:
        style_caption = styles['Caption']
    except KeyError:
        style_caption = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)

    font = style_caption.font
    font.name = FONT_NAME
    font.size = FONT_SIZES['caption']
    font.italic = True
    font.color.rgb = BRAND_COLORS['dark_gray']

    # --- Title Style ---
    style_title = styles['Title']
    font = style_title.font
    font.name = FONT_NAME
    font.size = Pt(18)
    font.bold = True
    font.color.rgb = BRAND_COLORS['black']


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def add_highlighted_text(
    paragraph,
    text: str,
    color: str = 'brand_blue',
    bold: bool = False
) -> None:
    """
    Add highlighted text to a paragraph using brand colors.

    Parameters
    ----------
    paragraph : Paragraph
        The paragraph to add text to
    text : str
        The text to add
    color : str
        Brand color name
    bold : bool
        Whether to make the text bold
    """
    run = paragraph.add_run(text)
    run.font.color.rgb = BRAND_COLORS.get(color, BRAND_COLORS['brand_blue'])
    run.font.bold = bold


def create_brand_table(
    doc: Document,
    data: list,
    col_widths: Optional[list] = None,
    header_bg: str = 'light_gray'
) -> 'Table':
    """
    Create a table with brand styling.

    Parameters
    ----------
    doc : Document
        The document to add the table to
    data : list
        2D list of table data (first row = headers)
    col_widths : list, optional
        Column widths in inches
    header_bg : str
        Background color for header row

    Returns
    -------
    Table
        The created table

    Examples
    --------
    >>> data = [['Name', 'Value'], ['A', '1'], ['B', '2']]
    >>> table = create_brand_table(doc, data)
    """
    rows = len(data)
    cols = len(data[0]) if data else 0

    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Apply column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    # Populate and style table
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_data)

            # Style the cell
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in paragraph.runs:
                run.font.name = FONT_NAME
                run.font.size = Pt(9)

            # Header row styling
            if i == 0:
                for run in paragraph.runs:
                    run.font.bold = True
                # Set background color
                set_cell_shading(cell, BRAND_HEX.get(header_bg, '#BDBDBD'))

    return table


def set_cell_shading(cell, hex_color: str) -> None:
    """
    Set the background color of a table cell.

    Parameters
    ----------
    cell : Cell
        The table cell
    hex_color : str
        Hex color string (e.g., '#BDBDBD')
    """
    # Remove # if present
    color = hex_color.lstrip('#')

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_figure_caption(
    doc: Document,
    caption_text: str,
    figure_number: Optional[int] = None
) -> None:
    """
    Add a figure caption with brand styling.

    Parameters
    ----------
    doc : Document
        The document
    caption_text : str
        The caption text
    figure_number : int, optional
        Figure number to prepend
    """
    if figure_number:
        full_caption = f"Figure {figure_number}. {caption_text}"
    else:
        full_caption = caption_text

    para = doc.add_paragraph(style='Caption')
    para.add_run(full_caption)


def add_table_caption(
    doc: Document,
    caption_text: str,
    table_number: Optional[int] = None
) -> None:
    """
    Add a table caption with brand styling.

    Parameters
    ----------
    doc : Document
        The document
    caption_text : str
        The caption text
    table_number : int, optional
        Table number to prepend
    """
    if table_number:
        full_caption = f"Table {table_number}. {caption_text}"
    else:
        full_caption = caption_text

    para = doc.add_paragraph(style='Caption')
    para.add_run(full_caption)


# =============================================================================
# EXAMPLE USAGE
# =============================================================================

def create_example_document():
    """Create an example document demonstrating brand styles."""
    doc = create_brand_document()

    # Title
    doc.add_heading("Oligon Scientific Report", level=0)

    # Introduction
    doc.add_heading("Introduction", level=1)
    para = doc.add_paragraph(
        "This document demonstrates the use of Oligon brand guidelines "
        "in Word document generation. All colors, fonts, and spacing "
        "follow the brand specification."
    )

    # Add highlighted text
    doc.add_heading("Key Findings", level=2)
    para = doc.add_paragraph("The ")
    add_highlighted_text(para, "treatment group", color='brand_blue', bold=True)
    para.add_run(" showed significant improvement compared to the control group (p < 0.05).")

    # Add table
    doc.add_heading("Results Summary", level=2)
    data = [
        ['Group', 'N', 'Mean', 'SD', 'p-value'],
        ['Control', '25', '4.2', '1.1', '-'],
        ['Treatment', '25', '6.8', '0.9', '<0.001'],
    ]
    create_brand_table(doc, data, col_widths=[1.5, 0.75, 1.0, 0.75, 1.0])
    add_table_caption(doc, "Summary statistics for treatment comparison.", table_number=1)

    doc.save("example_brand_document.docx")
    print("Example document created: example_brand_document.docx")


if __name__ == "__main__":
    create_example_document()
